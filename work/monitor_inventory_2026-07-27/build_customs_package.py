import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(r"C:\Users\AlejandroAcosta\Documents\ai-workstation\work\monitor_inventory_2026-07-27")
OUTPUT = Path(r"C:\Users\AlejandroAcosta\Documents\ai-workstation\outputs\monitor_inventory_2026-07-27")
TEMPLATE = Path(r"C:\Users\AlejandroAcosta\OneDrive - Interwork Office Solutions\Desktop\Blank Interwork BOL -.docx")
LOGO = BASE / "bol_template_assets" / "image1.png"
OUTFILE = OUTPUT / "project_7526_mexico_monitor_customs_information_package_interwork.docx"

NAVY = "767171"
BLUE = "C00000"
PALE_BLUE = "BFBFBF"
LIGHT_GRAY = "D9D9D9"
GREEN = "F2F2F2"
GOLD = "F2F2F2"
RED = "9B1C1C"
GRAY = "767171"

MODEL_MAP = {
    "IMG_6306.JPG": ("ThinkVision T22i-10", "A16215FT0"),
    "IMG_6326.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6375.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6376.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6383.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6387.JPG": ("ThinkVision T22i-10", "A16215FT0"),
    "IMG_6393.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6397.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6401.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6403.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6404.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6408.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6410.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6411.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6416.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6424.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6442.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6445.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6448.JPG": ("ThinkVision T2254pC", "T2254pC"),
    "IMG_6459.JPG": ("ThinkVision T22i-10", "A16215FT0"),
    "IMG_6463.JPG": ("ThinkVision T2254pC", "T2254pC"),
}

SERIAL_MAP = {
    "IMG_6326.JPG": "VNA1H63AK",
    "IMG_6383.JPG": "VNA21P7P",
    "IMG_6393.JPG": "VNA21POY",
    "IMG_6404.JPG": "VNA21N77",
    "IMG_6408.JPG": "VNA2PHF",
    "IMG_6448.JPG": "VNA21PD4",
}

MODEL_SPECS = [
    {
        "series": "ThinkVision T2254pC",
        "type": "T2254pC",
        "mtm": "60E1-MAR2-WW",
        "fru": "00PC018",
        "qty": 151,
        "description": "22-inch widescreen LED monitor; TN panel; 1680 × 1050; 16:10; VGA, HDMI 1.4 and DisplayPort 1.2; VESA 100 × 100 mm.",
        "origin": "China — confirmed for this inventory",
        "spec": "https://psref.lenovo.com/syspool/Sys/PDF/ThinkVision/ThinkVision_T2254p/ThinkVision_T2254p_Spec.PDF",
        "manual": "https://support.lenovo.com/ng/en/solutions/pd500342-thinkvision-t2254p-22inch-led-61ba-monitor-overview-and-service-parts",
        "photo": "IMG_6448.JPG",
    },
    {
        "series": "ThinkVision T22i-10",
        "type": "A16215FT0",
        "mtm": "61A9-MAR1-WW",
        "fru": "00PC160",
        "qty": 25,
        "description": "21.5-inch IPS monitor; 1920 × 1080 Full HD; 16:9; VGA, HDMI 1.4 and DisplayPort; four USB 3.0 ports; VESA 100 × 100 mm.",
        "origin": "China — confirmed for this inventory",
        "spec": "https://psref.lenovo.com/syspool/Sys/PDF/ThinkVision/ThinkVision_T22i_10/ThinkVision_T22i_Spec.pdf",
        "manual": "https://support.lenovo.com/gb/en/solutions/t22i_10/",
        "photo": "IMG_6251.JPG",
    },
    {
        "series": "ThinkVision T22i-20",
        "type": "D20215FT0",
        "mtm": "61FE-MAR6-WW",
        "fru": "5D10Y48252",
        "qty": 5,
        "description": "21.5-inch IPS monitor; 1920 × 1080 Full HD; 16:9; VGA, HDMI 1.4 and DisplayPort 1.2; USB hub; VESA 100 × 100 mm.",
        "origin": "China — confirmed for this inventory",
        "spec": "https://psref.lenovo.com/syspool/Sys/PDF/ThinkVision/ThinkVision_T22i_20/ThinkVision_T22i_20_Spec.PDF",
        "manual": "https://support.lenovo.com/sr/en/solutions/pd500410/",
        "photo": "IMG_6347.JPG",
    },
    {
        "series": "ThinkVision P24h-30",
        "type": "D22238QP0",
        "mtm": "63B3-GAR6-WW",
        "fru": "5D11J31074",
        "qty": 9,
        "description": "23.8-inch IPS monitor; 2560 × 1440 QHD; 16:9; USB-C docking with power delivery, HDMI 2.0, DisplayPort 1.4, DisplayPort out, USB hub and Ethernet.",
        "origin": "China — confirmed for this inventory",
        "spec": "https://psref.lenovo.com/syspool/Sys/PDF/ThinkVision/ThinkVision_P24h_30/ThinkVision_P24h_30_Spec.PDF",
        "manual": "https://support.lenovo.com/in/en/solutions/pd500607-thinkvision-p24h-30-lcd-monitor-overview",
        "photo": "IMG_6253.JPG",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)
            row.cells[idx]._tc.tcPr.tcW.set(qn("w:w"), str(int(width * 1440)))
            row.cells[idx]._tc.tcPr.tcW.set(qn("w:type"), "dxa")
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])


def font_run(run, size=10, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text="", size=11, bold=False, color="000000", after=6, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    font_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def create_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, text):
    p = doc.add_paragraph()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(doc._interwork_bullet_num_id))
    num_pr.extend([ilvl, num_id])
    p._p.get_or_add_pPr().append(num_pr)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    font_run(p.add_run(text), size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), BLUE)
        p_pr.append(shd)
        font_run(run, size=13, bold=True, color="FFFFFF")
        p.paragraph_format.left_indent = Inches(0.06)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    else:
        font_run(run, size=11.5, bold=True, color=BLUE)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    font_run(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def style_table_header(table):
    set_repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, NAVY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                font_run(run, size=9, bold=True, color="FFFFFF")
            paragraph.paragraph_format.space_after = Pt(0)


def add_label_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_GRAY)
        for run in cells[0].paragraphs[0].runs:
            font_run(run, size=10, bold=True, color=NAVY)
        for run in cells[1].paragraphs[0].runs:
            font_run(run, size=10)
    set_table_widths(table, [1.875, 4.625])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


records = json.loads((BASE / "reconciled_inventory.json").read_text(encoding="utf-8"))
for row in records:
    if row["photo_file"] in MODEL_MAP:
        row["series_name"], row["type_model_code"] = MODEL_MAP[row["photo_file"]]
    if row["photo_file"] in SERIAL_MAP:
        row["serial_number"] = SERIAL_MAP[row["photo_file"]]

records.sort(key=lambda row: row["photo_file"].upper())

doc = Document()
doc._interwork_bullet_num_id = create_bullet_numbering(doc)
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(0.31)
section.right_margin = Inches(0.28)
section.top_margin = Inches(0.19)
section.bottom_margin = Inches(0.01)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
):
    try:
        style = styles[style_name]
    except KeyError:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

footer = section.footer
add_page_number(footer.paragraphs[0])

logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
logo_p.paragraph_format.space_after = Pt(3)
logo_p.add_run().add_picture(str(LOGO), width=Inches(2.55))

title_band = doc.add_table(rows=1, cols=1)
title_band.style = "Table Grid"
set_table_widths(title_band, [6.5])
set_cell_shading(title_band.cell(0, 0), BLUE)
title_p = title_band.cell(0, 0).paragraphs[0]
font_run(title_p.add_run("CUSTOMS INFORMATION PACKAGE"), size=16, bold=True, color="FFFFFF")

add_text(doc, "Project 7526 - Used Lenovo Monitor Shipment", size=20, bold=True, color=NAVY, after=3)
add_text(doc, "Technical specifications, product identification, photographs, and serial-number annex", size=11.5, color=GRAY, after=9)

add_label_value_table(doc, [
    ("Prepared for", "Sunset Transportation / Mexican customs classification team"),
    ("Shipment", "One-time shipment from the United States to Mexico"),
    ("Inventory basis", "190 client-owned used Lenovo monitors"),
    ("Working unit value", "USD $75.00 per monitor"),
    ("Working total value", "USD $14,250.00"),
    ("Document date", "July 27, 2026"),
])

notice = doc.add_table(rows=1, cols=1)
notice.style = "Table Grid"
set_table_widths(notice, [6.5])
set_cell_shading(notice.cell(0, 0), GOLD)
p = notice.cell(0, 0).paragraphs[0]
font_run(p.add_run("FINAL SIGN-OFF REQUIRED: "), size=10, bold=True, color=RED)
font_run(p.add_run("Confirm the importer-of-record and final commercial-invoice parties before customs submission."), size=10, color=RED)

add_heading(doc, "1. Executive shipment summary", 1)
add_text(doc, "This package responds to the technical-information request for Mexican tariff classification and NOM review. All 190 physical units have an assigned model and serial value for working-ledger purposes.")

summary = doc.add_table(rows=1, cols=6)
summary.style = "Table Grid"
summary.rows[0].cells[0].text = "Manufacturer"
summary.rows[0].cells[1].text = "Series"
summary.rows[0].cells[2].text = "Type/Model"
summary.rows[0].cells[3].text = "MTM"
summary.rows[0].cells[4].text = "FRU"
summary.rows[0].cells[5].text = "Qty."
for spec in MODEL_SPECS:
    cells = summary.add_row().cells
    values = ["Lenovo", spec["series"], spec["type"], spec["mtm"], spec["fru"], str(spec["qty"])]
    for idx, value in enumerate(values):
        cells[idx].text = value
set_table_widths(summary, [0.72, 1.55, 1.10, 1.20, 1.15, 0.53])
style_table_header(summary)
for row in summary.rows[1:]:
    row.cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            font_run(run, size=8.5)

add_heading(doc, "2. Commercial product description", 1)
add_text(doc, "Used Lenovo ThinkVision flat-panel computer monitors, client-owned, non-hazardous, without batteries. Mixed models and production years. Intended for reuse as computer display equipment.")
add_bullet(doc, "Quantity: 190 units")
add_bullet(doc, "Working unit value: USD $75.00")
add_bullet(doc, "Working extended value: USD $14,250.00")
add_bullet(doc, "Shipment origin: United States")
add_bullet(doc, "Destination: Mexico")
add_bullet(doc, "Final commercial invoice seller/exporter, consignee and importer-of-record: pending confirmation")

add_heading(doc, "3. Model specifications and documentation", 1)
for spec in MODEL_SPECS:
    add_heading(doc, f'{spec["series"]} — {spec["type"]}', 2)
    add_label_value_table(doc, [
        ("Manufacturer", "Lenovo"),
        ("Model / part identifiers", f'{spec["series"]}; Type/Model {spec["type"]}; MTM {spec["mtm"]}; FRU {spec["fru"]}'),
        ("Quantity", str(spec["qty"])),
        ("Technical description", spec["description"]),
        ("Country of origin", spec["origin"]),
        ("Specification sheet", spec["spec"]),
        ("User manual / support", spec["manual"]),
    ])

add_heading(doc, "4. Representative product-label photographs", 1)
add_text(doc, "The photographs below show representative physical units from the inventory and provide model, part and serial-number evidence.")
for idx, spec in enumerate(MODEL_SPECS):
    photo_path = BASE / "monitors" / spec["photo"]
    add_heading(doc, f'{spec["series"]} — representative label ({spec["photo"]})', 2)
    if photo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(photo_path), width=Inches(6.1))
        p.paragraph_format.space_after = Pt(4)
    add_text(doc, f'Type/Model: {spec["type"]}  |  MTM: {spec["mtm"]}  |  FRU: {spec["fru"]}', size=9.5, color=GRAY, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    if idx < len(MODEL_SPECS) - 1:
        doc.add_page_break()

add_heading(doc, "5. Country-of-origin evidence status", 1)
origin_table = doc.add_table(rows=1, cols=3)
origin_table.style = "Table Grid"
origin_table.rows[0].cells[0].text = "Model"
origin_table.rows[0].cells[1].text = "Country of origin"
origin_table.rows[0].cells[2].text = "Evidence status"
for spec in MODEL_SPECS:
    cells = origin_table.add_row().cells
    cells[0].text = spec["series"]
    cells[1].text = "China"
    cells[2].text = "Confirmed for this inventory by project direction"
set_table_widths(origin_table, [1.65, 1.15, 3.70])
style_table_header(origin_table)

add_heading(doc, "6. Submission checklist", 1)
for item in (
    "Technical specification sheets: supplied by official Lenovo links for all four models.",
    "Manufacturer/model/part numbers: included in the model summary and specification sections.",
    "User manuals: supplied through official Lenovo support links.",
    "Product-label photographs: supplied for all four models.",
    "Complete serial-number list: Appendix A.",
    "Commercial description and working value: included; final invoice parties require confirmation.",
    "Country of origin: China for all four model groups.",
    "Mexican tariff classification, HTS determination and NOM applicability: to be completed by the customs broker/importer using this technical package.",
):
    add_bullet(doc, item)

doc.add_page_break()
add_heading(doc, "Appendix A — Complete 190-unit serial and model ledger", 1)
add_text(doc, "All serial values are treated as accepted for purposes of this package, per project direction. Source-photo filenames are retained for audit traceability.", size=9.5, color=GRAY)

ledger = doc.add_table(rows=1, cols=5)
ledger.style = "Table Grid"
for idx, label in enumerate(("No.", "Source Photo", "Series", "Type/Model", "Serial Number")):
    ledger.rows[0].cells[idx].text = label
for index, row in enumerate(records, start=1):
    cells = ledger.add_row().cells
    values = [
        str(index),
        row["photo_file"],
        row.get("series_name", ""),
        row.get("type_model_code", ""),
        row.get("serial_number", ""),
    ]
    for idx, value in enumerate(values):
        cells[idx].text = str(value or "")
set_table_widths(ledger, [0.42, 1.15, 2.03, 1.30, 1.60])
style_table_header(ledger)
for row in ledger.rows[1:]:
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                font_run(run, size=8)

add_text(doc, "End of package.", size=9, color=GRAY, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

OUTPUT.mkdir(parents=True, exist_ok=True)
doc.save(OUTFILE)
print(OUTFILE)
